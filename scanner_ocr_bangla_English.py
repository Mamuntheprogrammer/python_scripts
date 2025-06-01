# import os
# from pdf2image import convert_from_path
# import pytesseract
# from docx import Document
# from docx.shared import Pt
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement

# # Configuration
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# # poppler_path = r'C:\path\to\poppler\bin'  # Use if needed on Windows
# output_docx_path = "formatted_output.docx"
# font_name = "Siyam Rupali"
# font_size = 14

# def clean_text(text):
#     lines = [line.strip() for line in text.splitlines() if line.strip()]
#     return lines

# def set_bangla_font(run, font_name='Siyam Rupali'):
#     run.font.name = font_name
#     run.font.size = Pt(font_size)
#     rPr = run._element.rPr
#     rFonts = rPr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
#     rFonts.set(qn('w:ascii'), font_name)
#     rFonts.set(qn('w:hAnsi'), font_name)
#     rFonts.set(qn('w:cs'), font_name)
#     rPr.append(rFonts)

# def load_or_create_doc(path):
#     if os.path.exists(path):
#         print(f"Loading existing DOCX: {path}")
#         return Document(path)
#     else:
#         print(f"Creating new DOCX: {path}")
#         return Document()

# def process_page(image, page_number, doc):
#     doc.add_paragraph(f"--- Page {page_number} ---")

#     text = pytesseract.image_to_string(image, lang='ben')
#     lines = clean_text(text)

#     for line in lines:
#         para = doc.add_paragraph()
#         run = para.add_run(line)
#         set_bangla_font(run, font_name)

#     doc.add_page_break()

# def main(pdf_path, start_page=1):
#     # Convert only the needed page one-by-one
#     total_pages = len(convert_from_path(pdf_path, first_page=1, last_page=1))
#     print(f"Estimating total pages...")

#     page = start_page
#     while True:
#         try:
#             print(f"\nProcessing Page {page}...")

#             # Load single page only
#             images = convert_from_path(pdf_path, first_page=page, last_page=page, dpi=300)
#             image = images[0]

#             # Load existing doc or create new
#             doc = load_or_create_doc(output_docx_path)

#             # OCR & add to doc
#             process_page(image, page, doc)

#             # Save incrementally
#             doc.save(output_docx_path)
#             print(f"Page {page} processed and saved.")

#             page += 1
#         except Exception as e:
#             print(f"\nStopped at page {page}. Reason: {e}")
#             print("You can resume later by setting start_page=", page)
#             break

# if __name__ == "__main__":
#     # Set your input PDF and resume point
#     pdf_path = "ict.pdf"  # 🔁 Change this
#     start_page = 1              # 🔁 Change this if resuming

#     main(pdf_path, start_page=start_page)


import os
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

# Tesseract OCR config
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Constants
output_docx_path = "Output.docx"
font_name = "Siyam Rupali"
font_size = 11  # Standard Word font size

def set_bangla_font(run, font_name):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    rPr = run._element.rPr
    rFonts = rPr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def create_or_load_doc(path):
    if os.path.exists(path):
        print(f"Loading existing DOCX: {path}")
        doc = Document(path)
    else:
        print(f"Creating new DOCX: {path}")
        doc = Document()

        # Set page size and orientation for first section
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc

def process_page(image, page_number, doc):
    # Page heading (optional)
    heading = doc.add_paragraph()
    run = heading.add_run(f"পৃষ্ঠা  Page - {page_number}")
    run.bold = True
    set_bangla_font(run, font_name)

    # Extract and clean text (one paragraph, no gaps)
    raw_text = pytesseract.image_to_string(image, lang='ben+eng')
    cleaned_text = ' '.join(line.strip() for line in raw_text.splitlines() if line.strip())

    if cleaned_text:
        para = doc.add_paragraph()
        run = para.add_run(cleaned_text)
        set_bangla_font(run, font_name)

    doc.add_page_break()

def main(pdf_path, start_page=1):
    print("Estimating total pages...")
    # Check how many pages exist (first page call just to get image count)
    total_images = convert_from_path(pdf_path, first_page=1, last_page=1)
    total_pages = len(total_images)

    page = start_page
    while True:
        try:
            print(f"\nProcessing Page {page}...")

            # Load single page
            images = convert_from_path(pdf_path, first_page=page, last_page=page, dpi=300)
            image = images[0]

            # Load or create doc
            doc = create_or_load_doc(output_docx_path)

            # OCR and write to doc
            process_page(image, page, doc)

            # Save after each page
            doc.save(output_docx_path)
            print(f"Page {page} saved to DOCX.")

            page += 1
        except Exception as e:
            print(f"\nStopped at page {page}. Reason: {e}")
            print("To resume, set start_page =", page)
            break

if __name__ == "__main__":
    # Change your file and resume page here
    pdf_path = "ict.pdf"
    start_page = 1
    main(pdf_path, start_page)
